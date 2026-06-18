import streamlit as st
import google.generativeai as genai
import sqlite3
import json
import os
import re
import random
from io import BytesIO
from datetime import datetime
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px

# Try importing parsing packages
try:
    import pypdf
    PDF_PARSING_AVAILABLE = True
except ImportError:
    PDF_PARSING_AVAILABLE = False

try:
    import docx
    DOCX_PARSING_AVAILABLE = True
except ImportError:
    DOCX_PARSING_AVAILABLE = False

# ── Database Setup ─────────────────────────────────────────
DB_PATH = "resume_analyzer.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            timestamp TEXT,
            score INTEGER,
            target_role TEXT,
            analysis_json TEXT
        )
    """)
    conn.commit()
    conn.close()

def save_analysis(filename, score, target_role, analysis_dict):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    analysis_json = json.dumps(analysis_dict)
    cursor.execute("""
        INSERT INTO history (filename, timestamp, score, target_role, analysis_json)
        VALUES (?, ?, ?, ?, ?)
    """, (filename, timestamp, score, target_role, analysis_json))
    conn.commit()
    conn.close()

def get_history():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT id, filename, timestamp, score, target_role, analysis_json FROM history ORDER BY id DESC")
    rows = cursor.fetchall()
    conn.close()
    return rows

def clear_history():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM history")
    conn.commit()
    conn.close()

# Initialize database
init_db()

# ── Text Parsing Helpers ────────────────────────────────────
def extract_text_from_pdf(file_bytes):
    if not PDF_PARSING_AVAILABLE:
        return "Error: pypdf library is not installed."
    try:
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"

def extract_text_from_docx(file_bytes):
    if not DOCX_PARSING_AVAILABLE:
        return "Error: python-docx library is not installed."
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"

# ── Offline Heuristic Analyzer (Fallback/Demo Mode) ──────────
SKILLS_LIBRARY = {
    "languages": ["python", "javascript", "typescript", "java", "c++", "c#", "ruby", "go", "rust", "php", "html", "css", "sql", "kotlin", "swift"],
    "frameworks": ["react", "angular", "vue", "next.js", "django", "flask", "fastapi", "spring boot", "express", "rails", "laravel", "flutter", "react native"],
    "databases": ["postgresql", "mysql", "mongodb", "sqlite", "redis", "oracle", "sql server", "cassandra", "dynamodb"],
    "devops_cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "git", "ci/cd", "jenkins", "terraform", "ansible"],
    "data_science": ["pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "matplotlib", "seaborn", "power bi", "tableau", "spark", "hadoop"]
}

def analyze_resume_heuristically(text, target_role, target_desc=""):
    text_lower = text.lower()
    desc_lower = target_desc.lower() if target_desc else ""
    role_lower = target_role.lower()
    
    # 1. Extract skills found in resume
    found_skills = []
    for category, skills in SKILLS_LIBRARY.items():
        for skill in skills:
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found_skills.append(skill.title())
                
    # 2. Extract potential keywords from target description / title
    target_keywords = []
    for category, skills in SKILLS_LIBRARY.items():
        for skill in skills:
            if re.search(r'\b' + re.escape(skill) + r'\b', desc_lower) or re.search(r'\b' + re.escape(skill) + r'\b', role_lower):
                target_keywords.append(skill.title())
                
    # If no target keywords found, pick some standard ones based on target role keyword matching
    if not target_keywords:
        if "data" in role_lower:
            target_keywords = ["Python", "SQL", "Pandas", "Power BI", "Tableau", "AWS"]
        elif "web" in role_lower or "frontend" in role_lower or "react" in role_lower:
            target_keywords = ["JavaScript", "TypeScript", "React", "HTML", "CSS", "Git"]
        elif "devops" in role_lower or "cloud" in role_lower:
            target_keywords = ["AWS", "Docker", "Kubernetes", "CI/CD", "Linux", "Git"]
        else:
            target_keywords = ["Python", "SQL", "Git", "Docker", "RESTful APIs"]
            
    # 3. Detect missing skills
    missing_skills = [skill for skill in target_keywords if skill not in found_skills]
    if not missing_skills:
        # Guarantee at least some learning suggestions for coach mode
        missing_skills = [s for s in ["Kubernetes", "AWS Developer", "System Design", "CI/CD Pipelines"] if s not in found_skills][:3]

    # 4. Check for contact details
    has_email = "@" in text_lower
    has_phone = bool(re.search(r'\b\d{10}\b|\b\d{3}[-.\s]??\d{3}[-.\s]??\d{4}\b', text_lower))
    has_linkedin = "linkedin" in text_lower
    
    # 5. Check formatting structures
    has_experience_section = any(k in text_lower for k in ["experience", "employment", "work history", "history"])
    has_education_section = any(k in text_lower for k in ["education", "academic", "university", "college", "degree"])
    has_projects_section = any(k in text_lower for k in ["projects", "personal projects", "portfolio"])
    
    # 6. Calculate subscores
    formatting_score = 60
    if has_experience_section: formatting_score += 15
    if has_education_section: formatting_score += 15
    if has_projects_section: formatting_score += 10
    
    skills_score = min(50 + len(found_skills) * 4, 100)
    
    experience_score = 65
    if has_experience_section:
        experience_score += 15
        if len(text_lower.split()) > 400:
            experience_score += 15
    experience_score = min(experience_score, 100)
    
    # Keyword match score
    if target_keywords:
        matched = [k for k in target_keywords if k in found_skills]
        keyword_score = int((len(matched) / len(target_keywords)) * 100) if target_keywords else 75
        keyword_score = max(50, min(keyword_score, 100))
    else:
        keyword_score = 70
        
    projects_score = 80 if has_projects_section else 50
    education_score = 85 if has_education_section else 55
    
    # Heuristics adjustment to simulate version changes based on file name identifiers
    # If file name contains V2 or V3, let's bump the score to show progress over time
    score_modifier = 0
    
    # ATS calculation
    ats_score = int((formatting_score + skills_score + experience_score + keyword_score + projects_score + education_score) / 6)
    
    # Sample rewrites and feedback
    strengths = [
        "Core foundational sections (Experience, Education) are clearly present." if has_experience_section and has_education_section else "Detailed description of technical activities.",
        f"Extracted {len(found_skills)} valid professional skills directly from your resume text.",
        "Clear technical projects description." if has_projects_section else "Easy-to-read textual layout format."
    ]
    
    weaknesses = []
    if not has_projects_section:
        weaknesses.append("Missing a dedicated 'Projects' section to showcase practical applications.")
    if not has_linkedin:
        weaknesses.append("No LinkedIn profile URL detected; recruiters value social profiles.")
    if len(missing_skills) > 0:
        weaknesses.append(f"Missing core keywords matching the target role: {', '.join(missing_skills[:3])}.")
    if not weaknesses:
        weaknesses = ["Could expand on impact metrics (e.g., %, $) in your work achievements.", "Skills section could be categorized for easier scanning."]
        
    improvements = [
        f"Incorporate missing skills: {', '.join(missing_skills)} to improve your search matching.",
        "Add quantifiable metrics to your experience bullet points (e.g., 'Reduced loading time by 20%').",
        "Add a professional profile summary at the top aligned with target role."
    ]
    
    suggested_roles = [target_role]
    if "data" in role_lower:
        suggested_roles += ["Data Engineer", "Business Intelligence Analyst", "Analytics Engineer"]
    elif "software" in role_lower or "developer" in role_lower:
        suggested_roles += ["Full Stack Developer", "Software Engineer II", "Systems Architect"]
    elif "cloud" in role_lower or "devops" in role_lower:
        suggested_roles += ["Site Reliability Engineer (SRE)", "Infrastructure Engineer", "Cloud Solutions Architect"]
    else:
        suggested_roles += ["Associate Project Manager", "Technical Consultant", "Operations Lead"]

    # Generate rewrites based on found skills
    skills_str = ", ".join(found_skills) if found_skills else "Python, SQL, Git, RESTful APIs"
    
    result = {
        "ats_score": ats_score,
        "score_breakdown": {
            "formatting": formatting_score,
            "skills": skills_score,
            "experience": experience_score,
            "keywords": keyword_score,
            "projects": projects_score,
            "education": education_score
        },
        "score_explanation": f"Your resume scored a {ats_score}/100. You have a solid academic foundation and format. However, to increase compatibility with modern applicant tracking systems for '{target_role}', you need to explicitly incorporate missing keywords (like {', '.join(missing_skills[:3])}) and replace passive sentences with impact-driven action statements.",
        "skills_extracted": found_skills if found_skills else ["Python", "SQL", "Git", "Excel"],
        "missing_skills": missing_skills,
        "missing_keywords": [s.lower() for s in missing_skills],
        "missing_certifications": [f"Certified {target_role} Professional", "AWS Certified Cloud Practitioner", "Professional Scrum Master I"][:2],
        "strengths": strengths,
        "weaknesses": weaknesses,
        "improvements": improvements,
        "interview_readiness": "Moderate. Your resume displays technical knowledge but needs keyword alignment and impact formatting to reliably pass applicant screening systems.",
        "rewritten_summary": f"Professional and results-driven specialist with hands-on expertise in {skills_str}. Proven track record of developing functional systems, designing efficient data structures, and solving complex problems. Adept at collaborating in cross-functional teams to deploy high-quality software aligned with '{target_role}' target parameters.",
        "rewritten_skills": f"{skills_str}, Agile Methodologies, Solution Design, Analytical Troubleshooting",
        "rewritten_experience": [
            f"Optimized system processes and tools related to {target_role} functions, improving workflow efficiency by 22%.",
            f"Designed and deployed modular solutions utilizing {found_skills[0] if found_skills else 'core tools'}, decreasing processing delays by 15%.",
            "Collaborated with cross-functional stakeholders to define technical specifications and verify deployment standards."
        ],
        "rewritten_projects": [
            f"Automated System Architect: Designed a custom automation engine using {found_skills[0] if found_skills else 'core scripting'}, replacing manual pipelines and saving 10+ engineering hours weekly.",
            "Integrated Data Analytics Hub: Built a personal analytics dashboard utilizing SQL and visualization charts to model key dataset patterns."
        ],
        "career_suggestions": {
            "suitable_roles": suggested_roles[:3],
            "learning_roadmap": [
                f"Week 1: Focus on mastering {missing_skills[0] if len(missing_skills) > 0 else 'advanced scripting'} and updating resume summary.",
                f"Week 2: Build a sandbox project integrating {missing_skills[1] if len(missing_skills) > 1 else 'cloud containers'} with databases.",
                f"Week 3: Deepen understanding of automated unit testing and system design.",
                f"Week 4: Execute practice interviews using mock questions and apply for target roles."
            ],
            "skills_to_learn_next": missing_skills[:3],
            "certifications_to_pursue": [f"Certified {target_role} Professional", "Google Professional Data/DevOps Certificate", "CompTIA Security+"][:2]
        },
        "interview_prep": {
            "technical_questions": [
                f"What are the best practices for structuring code/systems when developing solutions for '{target_role}'?",
                f"How would you integrate {found_skills[0] if found_skills else 'databases'} to ensure data consistency and minimize query latency?",
                "Walk me through a complex technical issue you encountered in a previous project and how you debugged it."
            ],
            "hr_questions": [
                f"Why are you interested in pursuing a career as a '{target_role}', and how do your skills prepare you for it?",
                "Describe a situation where you had to work under a tight deadline with incomplete requirements. How did you proceed?",
                "How do you keep your skills updated in a rapidly changing technical landscape?"
            ]
        },
        "roadmap_30_day": {
            "week_1": "Update resume structure, inject missing keywords, and set up your portfolio page.",
            "week_2": f"Implement a mini project incorporating {missing_skills[0] if len(missing_skills) > 0 else 'containerization'} to showcase hands-on capability.",
            "week_3": "Review core system design principles and start certification modules.",
            "week_4": "Begin active applications, leverage LinkedIn referrals, and perform technical mock practice."
        }
    }
    return result

# ── AI API Analysis ──────────────────────────────────────────
def analyze_resume_with_gemini(text, target_role, target_desc="", api_key=""):
    if not api_key:
        return {"error": "API Key is missing."}
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""You are a professional recruiting manager, ATS scanner expert, and career coach.
Analyze the following resume text against the target role: "{target_role}" and optional job description details: "{target_desc}".

Resume Text:
\"\"\"{text}\"\"\"

Generate a comprehensive analysis and return the result strictly as a JSON object matching this schema structure. Do not include any wrapper backticks like ```json or markdown text, return raw JSON string only:
{{
  "ats_score": 78,
  "score_breakdown": {{
    "formatting": 85,
    "skills": 70,
    "experience": 75,
    "keywords": 65,
    "projects": 80,
    "education": 90
  }},
  "score_explanation": "Provide a detailed overview of why they scored this. Mention key gaps.",
  "skills_extracted": ["Skill1", "Skill2"],
  "missing_skills": ["SkillA", "SkillB"],
  "missing_keywords": ["keyword1", "keyword2"],
  "missing_certifications": ["CertA"],
  "strengths": ["Strength1", "Strength2"],
  "weaknesses": ["Weakness1", "Weakness2"],
  "improvements": ["Improvement1", "Improvement2"],
  "interview_readiness": "A short summary paragraph of their interview readiness level.",
  "rewritten_summary": "An ATS-optimized professional resume summary statement for this candidate.",
  "rewritten_skills": "A formatted skills list, categorized (e.g. Languages: Python, JavaScript; Frameworks: React, Django).",
  "rewritten_experience": ["Bullet 1 with impact metrics", "Bullet 2 with impact metrics"],
  "rewritten_projects": ["Project bullet 1", "Project bullet 2"],
  "career_suggestions": {{
    "suitable_roles": ["Role1", "Role2"],
    "learning_roadmap": ["Step 1 description", "Step 2 description"],
    "skills_to_learn_next": ["SkillX", "SkillY"],
    "certifications_to_pursue": ["CertX", "CertY"]
  }},
  "interview_prep": {{
    "technical_questions": ["Technical Question 1", "Technical Question 2", "Technical Question 3"],
    "hr_questions": ["HR Question 1", "HR Question 2", "HR Question 3"]
  }},
  "roadmap_30_day": {{
    "week_1": "Goals for week 1",
    "week_2": "Goals for week 2",
    "week_3": "Goals for week 3",
    "week_4": "Goals for week 4"
  }}
}}
"""
        response = model.generate_content(prompt)
        raw_text = response.text.strip()
        
        # Clean JSON wrappers if model returns them
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]
        raw_text = raw_text.strip()
        
        return json.loads(raw_text)
    except Exception as e:
        return {"error": f"Failed to call Gemini API: {str(e)}"}

# ── Streamlit UI Setup ───────────────────────────────────────
st.set_page_config(
    page_title="AI Resume Coach & Analyzer (Multi-Resume Edition)",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium Style Tokens & CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="css"], .stMarkdown {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }
    
    .hero-container {
        padding: 3rem 2rem;
        background: linear-gradient(135deg, #090d16 0%, #15103a 50%, #290838 100%);
        color: white;
        border-radius: 24px;
        margin-bottom: 2.5rem;
        text-align: center;
        border: 1px solid rgba(255, 255, 255, 0.08);
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.25);
    }
    
    .hero-container h1 {
        font-weight: 700;
        font-size: 3rem;
        letter-spacing: -1px;
        background: linear-gradient(90deg, #38bdf8 0%, #a78bfa 50%, #f472b6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.5rem;
    }
    
    .hero-container p {
        font-size: 1.2rem;
        font-weight: 300;
        opacity: 0.9;
        margin-bottom: 0;
    }
    
    .metric-card {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 16px;
        padding: 1.5rem;
        text-align: center;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        margin: 0.5rem 0;
    }
    
    .metric-title {
        font-size: 0.85rem;
        color: #64748b;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-weight: 600;
        margin-bottom: 0.25rem;
    }
    
    .metric-value {
        font-size: 2.2rem;
        font-weight: 700;
        color: #0f172a;
    }
    
    .feature-box {
        background: rgba(248, 250, 252, 0.8);
        border: 1px solid #e2e8f0;
        border-radius: 16px;
        padding: 1.8rem;
        border-left: 5px solid #6366f1;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.02);
        margin: 1rem 0;
    }
    
    .interactive-bullet {
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
        transition: transform 0.2s ease;
    }
    
    .interactive-bullet:hover {
        transform: translateX(4px);
        border-color: #cbd5e1;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 48px;
        background-color: #f1f5f9;
        border-radius: 12px;
        font-weight: 500;
        padding: 0 20px;
        border: 1px solid transparent;
        transition: all 0.2s ease;
        color: #475569;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #e2e8f0;
        color: #0f172a;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%) !important;
        color: white !important;
        box-shadow: 0 4px 12px rgba(99, 102, 241, 0.3);
    }
    
    .roadmap-week {
        border-left: 3px solid #8b5cf6;
        padding-left: 1.5rem;
        margin-left: 0.5rem;
        position: relative;
        margin-bottom: 1.5rem;
    }
    
    .roadmap-week::before {
        content: '';
        width: 12px;
        height: 12px;
        background: #8b5cf6;
        border-radius: 50%;
        position: absolute;
        left: -8px;
        top: 4px;
        box-shadow: 0 0 0 4px rgba(139, 92, 246, 0.25);
    }
</style>
""", unsafe_allow_html=True)

# ── Sidebar ─────────────────────────────────────────────────
with st.sidebar:
    st.image("https://img.icons8.com/color/96/resume.png", width=64)
    st.markdown("### 💼 AI Resume & Career Coach")
    st.markdown("Optimize your resumes, compare versions side-by-side, and track career development progress.")
    
    st.divider()
    
    st.markdown("#### 🔑 Configuration")
    api_provider = st.selectbox("API Engine:", ["Gemini Demo Mode (Simulated)", "Google Gemini Live"], key="api_select")
    
    gemini_key = ""
    if api_provider == "Google Gemini Live":
        gemini_key = st.text_input(
            "Enter Gemini API Key:",
            type="password",
            help="Create a key at Google AI Studio to unlock advanced resume analysis.",
            key="gemini_key_input"
        )
        if not gemini_key:
            st.info("💡 Input your API Key to run live scans, or switch to Demo Mode to test instantly.")
    else:
        st.success("🤖 Demo Mode active! Locally scanning resumes.")

    st.divider()
    
    st.markdown("#### ⚙️ Target Job Profile")
    target_role = st.text_input("Target Job Title:", placeholder="e.g. Software Engineer, Data Analyst", value="Software Engineer")
    target_desc = st.text_area("Job Description / Keywords:", placeholder="Paste job descriptions here to enable keyword scanning...", height=120)
    
    st.divider()
    
    st.markdown("#### 📦 Module Diagnostics")
    if PDF_PARSING_AVAILABLE:
        st.markdown("✅ **PDF Extractor**: Active")
    else:
        st.markdown("⚠️ **PDF Extractor**: Unavailable (pypdf missing)")
        
    if DOCX_PARSING_AVAILABLE:
        st.markdown("✅ **Word Extractor**: Active")
    else:
        st.markdown("⚠️ **Word Extractor**: Unavailable (docx missing)")

# ── Landing Header ──────────────────────────────────────────
st.markdown("""
<div class="hero-container">
    <h1>AI Multi-Resume Analyzer & Career Coach</h1>
    <p>Upload multiple resumes side-by-side to compare, rank, and track improvements over time</p>
</div>
""", unsafe_allow_html=True)

# Initialize Session State
if "analyses" not in st.session_state:
    st.session_state.analyses = {}

# ── Multi-Upload Section ────────────────────────────────────
st.subheader("📤 Step 1: Upload Your Resumes")
uploaded_files = st.file_uploader(
    "Upload one or more PDF/Word DOCX resumes:", 
    type=["pdf", "docx"], 
    accept_multiple_files=True
)

if uploaded_files:
    # Build list of filenames currently uploaded
    uploaded_filenames = [f.name for f in uploaded_files]
    
    # Remove files from session state that are no longer uploaded
    for key in list(st.session_state.analyses.keys()):
        if key not in uploaded_filenames:
            del st.session_state.analyses[key]
            
    # Trigger button to process missing resumes
    unprocessed_files = [f for f in uploaded_files if f.name not in st.session_state.analyses]
    
    if unprocessed_files:
        st.warning(f"💡 You have {len(unprocessed_files)} new resume(s) to analyze.")
        col_btn_1, col_btn_2 = st.columns([1, 4])
        with col_btn_1:
            analyze_clicked = st.button("🚀 Process Resumes", use_container_width=True, type="primary")
            
        if analyze_clicked:
            for f in unprocessed_files:
                with st.spinner(f"Analyzing {f.name}... 🤖"):
                    f.seek(0)
                    file_bytes = f.read()
                    if f.name.endswith(".pdf"):
                        extracted_text = extract_text_from_pdf(file_bytes)
                    elif f.name.endswith(".docx"):
                        extracted_text = extract_text_from_docx(file_bytes)
                    else:
                        extracted_text = "Unsupported format."
                        
                    if extracted_text and "Error" not in extracted_text:
                        if api_provider == "Google Gemini Live":
                            if not gemini_key:
                                st.error("Please enter a valid Gemini API key in the sidebar.")
                                break
                            result = analyze_resume_with_gemini(extracted_text, target_role, target_desc, gemini_key)
                            if "error" in result:
                                st.error(f"Error on {f.name}: {result['error']}")
                            else:
                                st.session_state.analyses[f.name] = result
                                save_analysis(f.name, result["ats_score"], target_role, result)
                        else:
                            # Heuristic Analyzer
                            result = analyze_resume_heuristically(extracted_text, target_role, target_desc)
                            
                            # Artificial score adjustments to make multiple files look distinct in Demo Mode
                            if "v2" in f.name.lower() or "ver2" in f.name.lower() or "improved" in f.name.lower():
                                result["ats_score"] = min(result["ats_score"] + 8, 98)
                                result["score_breakdown"]["keywords"] = min(result["score_breakdown"]["keywords"] + 12, 100)
                                result["score_breakdown"]["formatting"] = min(result["score_breakdown"]["formatting"] + 5, 100)
                                result["score_explanation"] += " (V2 improvement adjustments loaded)"
                            elif "v3" in f.name.lower() or "ver3" in f.name.lower() or "final" in f.name.lower():
                                result["ats_score"] = min(result["ats_score"] + 14, 99)
                                result["score_breakdown"]["keywords"] = min(result["score_breakdown"]["keywords"] + 20, 100)
                                result["score_breakdown"]["formatting"] = min(result["score_breakdown"]["formatting"] + 10, 100)
                                result["score_explanation"] += " (V3 optimization adjustments loaded)"
                                
                            st.session_state.analyses[f.name] = result
                            save_analysis(f.name, result["ats_score"], target_role, result)
            st.rerun()

# ── Tabs Configuration ──────────────────────────────────────
st.markdown("<div style='height: 20px;'></div>", unsafe_allow_html=True)

tab_comparison, tab_dashboard, tab_details, tab_keywords, tab_rewriter, tab_coach, tab_prep, tab_history = st.tabs([
    "📂 Resume Comparison Panel",
    "📊 Individual Insights",
    "📈 ATS Score Details",
    "🔍 Keyword Matching",
    "✨ Resume Rewriter",
    "🧠 Career Roadmaps",
    "💬 Interview Practice",
    "📜 Analysis History"
])

# ── Global Comparison Variables ──────────────────────────────
has_data = len(st.session_state.analyses) > 0

# ── Tab 0: Comparison Panel ─────────────────────────────────
with tab_comparison:
    if not has_data:
        st.info("💡 Please upload resumes and click 'Process Resumes' to view comparisons.")
    else:
        st.markdown("### 📂 Resume Comparison Dashboard")
        st.markdown(f"Comparing uploaded resumes against target job profile: **{target_role}**")
        
        # Build comparison table data
        rows_list = []
        best_overall = None
        highest_score = -1
        
        for name, data in st.session_state.analyses.items():
            score = data["ats_score"]
            breakdown = data["score_breakdown"]
            rows_list.append({
                "Resume Name": name,
                "ATS Score": score,
                "Keyword Match": breakdown["keywords"],
                "Formatting": breakdown["formatting"],
                "Experience": breakdown["experience"],
                "Projects": breakdown["projects"],
                "Education": breakdown["education"]
            })
            if score > highest_score:
                highest_score = score
                best_overall = name
                
        df_comp = pd.DataFrame(rows_list)
        # Sort by ATS Score descending
        df_comp = df_comp.sort_values(by="ATS Score", ascending=False).reset_index(drop=True)
        
        # Add Overall Rank column
        df_comp.insert(0, "Overall Rank", range(1, len(df_comp) + 1))
        
        # Display comparison table
        st.markdown("#### Comparative Standings")
        st.dataframe(df_comp, use_container_width=True, hide_index=True)
        
        # Export Comparison Data
        csv_data = df_comp.to_csv(index=False).encode('utf-8')
        st.download_button(
            "⬇️ Export Comparison Data (CSV)",
            csv_data,
            "resume_comparison_matrix.csv",
            "text/csv",
            key="export_csv"
        )
        
        st.divider()
        
        # Best Resume Detection
        col_b1, col_b2 = st.columns(2)
        with col_b1:
            st.markdown(f"""
            <div style="background: #ecfdf5; border: 2px solid #10b981; border-radius: 16px; padding: 1.8rem; margin: 1rem 0;">
                <h3 style="color: #047857; margin-bottom: 0.5rem;">🏆 Best Overall Resume</h3>
                <h4 style="color: #065f46; font-size: 1.25rem;">{best_overall}</h4>
                <p style="color: #065f46; font-size: 0.95rem; margin-top: 0.5rem; line-height: 1.5;">
                    Selected due to scoring the highest compatibility match of <b>{highest_score}%</b> against your target parameters. This resume shows the strongest alignment of technical skills and keyword density.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
        with col_b2:
            st.markdown("#### 🥇 Category Winners")
            # Find category maxes
            best_fmt = df_comp.loc[df_comp["Formatting"].idxmax()]["Resume Name"]
            best_kw = df_comp.loc[df_comp["Keyword Match"].idxmax()]["Resume Name"]
            best_exp = df_comp.loc[df_comp["Experience"].idxmax()]["Resume Name"]
            
            st.markdown(f"✨ **Best Formatting**: `{best_fmt}` ({df_comp['Formatting'].max()}%)")
            st.markdown(f"🔑 **Best Keyword Match**: `{best_kw}` ({df_comp['Keyword Match'].max()}%)")
            st.markdown(f"💼 **Best Experience**: `{best_exp}` ({df_comp['Experience'].max()}%)")
            
        st.divider()
        
        # Visual Comparison Charts
        st.markdown("### 📊 Comparative Metrics Visualizations")
        col_c1, col_c2 = st.columns(2)
        
        with col_c1:
            st.markdown("#### Score Breakdown Comparison")
            # Grouped bar chart
            fig_bar = go.Figure()
            for idx, r_row in df_comp.iterrows():
                fig_bar.add_trace(go.Bar(
                    name=r_row["Resume Name"],
                    x=["ATS Score", "Keywords", "Formatting", "Experience"],
                    y=[r_row["ATS Score"], r_row["Keyword Match"], r_row["Formatting"], r_row["Experience"]]
                ))
            fig_bar.update_layout(
                barmode='group',
                yaxis=dict(range=[0, 100]),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(l=20, r=20, t=10, b=10)
            )
            st.plotly_chart(fig_bar, use_container_width=True)
            
        with col_c2:
            st.markdown("#### Version Improvement Tracking")
            # If multiple files exist, show progress line chart
            df_line = df_comp.copy()
            # Order by score or let user view progress
            fig_line = px.line(
                df_line, 
                x="Resume Name", 
                y="ATS Score", 
                markers=True, 
                title="ATS Score Progress Across Versions"
            )
            fig_line.update_layout(
                yaxis=dict(range=[0, 100]),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(l=20, r=20, t=10, b=10)
            )
            st.plotly_chart(fig_line, use_container_width=True)
            
        st.divider()
        
        # Version tracking details and improvements comparison
        st.markdown("### 🔄 Version Growth Details")
        st.markdown("Compare the specific changes and improvements across all analyzed files:")
        
        # If there are at least two resumes, let the user select V1 and V2 to compare
        if len(df_comp) >= 2:
            col_v1, col_v2 = st.columns(2)
            with col_v1:
                v1_select = st.selectbox("Compare Base Version (V1):", df_comp["Resume Name"].tolist(), index=len(df_comp)-1)
            with col_v2:
                v2_select = st.selectbox("Compare Target Version (V2):", df_comp["Resume Name"].tolist(), index=0)
                
            if v1_select != v2_select:
                v1_data = st.session_state.analyses[v1_select]
                v2_data = st.session_state.analyses[v2_select]
                
                score_diff = v2_data["ats_score"] - v1_data["ats_score"]
                skills_v1 = set(v1_data["skills_extracted"])
                skills_v2 = set(v2_data["skills_extracted"])
                added_skills = list(skills_v2 - skills_v1)
                
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    st.metric("ATS Score Growth", f"{v2_data['ats_score']}%", f"{score_diff:+}% improvement")
                with col_d2:
                    st.markdown("🎒 **Newly Extracted Skills**:")
                    if added_skills:
                        st.markdown(", ".join([f"`{s}`" for s in added_skills]))
                    else:
                        st.write("No new skills detected between these versions.")
        else:
            st.info("Upload 2 or more resumes to track version comparisons automatically.")

# ── Single Resume Selector Dropdown ─────────────────────────
selected_file = None
if has_data:
    st.divider()
    selected_file = st.selectbox(
        "🔎 Select Resume to Inspect Detailed Tab Content:", 
        list(st.session_state.analyses.keys()),
        help="Change this dropdown to update individual feedback details below."
    )

# ── Tab 1: Dashboard ────────────────────────────────────────
with tab_dashboard:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to view dashboard.")
    else:
        res = st.session_state.analyses[selected_file]
        
        # Display Metrics Cards
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        with col_m1:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-title">ATS Score</div>
                <div class="metric-value" style="color: #4f46e5;">{res['ats_score']}%</div>
            </div>
            """, unsafe_allow_html=True)
        with col_m2:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-title">Keywords Match</div>
                <div class="metric-value" style="color: #06b6d4;">{res['score_breakdown']['keywords']}%</div>
            </div>
            """, unsafe_allow_html=True)
        with col_m3:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-title">Formatting Rating</div>
                <div class="metric-value" style="color: #10b981;">{res['score_breakdown']['formatting']}%</div>
            </div>
            """, unsafe_allow_html=True)
        with col_m4:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-title">Experience Rating</div>
                <div class="metric-value" style="color: #f59e0b;">{res['score_breakdown']['experience']}%</div>
            </div>
            """, unsafe_allow_html=True)
            
        st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
        
        col_left, col_right = st.columns(2)
        with col_left:
            st.markdown("### 🏆 Core Strengths")
            for strength in res["strengths"]:
                st.markdown(f"<div class='interactive-bullet'>✅ {strength}</div>", unsafe_allow_html=True)
                
            st.markdown("### ⚠️ Key Gaps & Weaknesses")
            for weakness in res["weaknesses"]:
                st.markdown(f"<div class='interactive-bullet' style='border-left: 5px solid #ef4444;'>❌ {weakness}</div>", unsafe_allow_html=True)
                
        with col_right:
            st.markdown("### 📊 Skills & Category Weights")
            categories = list(res["score_breakdown"].keys())
            scores = list(res["score_breakdown"].values())
            
            fig = go.Figure(data=[
                go.Bar(
                    x=scores,
                    y=[cat.title() for cat in categories],
                    orientation='h',
                    marker=dict(
                        color=scores,
                        colorscale='Sunset',
                        line=dict(color='#cbd5e1', width=1)
                    )
                )
            ])
            fig.update_layout(
                margin=dict(l=20, r=20, t=10, b=10),
                height=260,
                xaxis=dict(range=[0, 100]),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)'
            )
            st.plotly_chart(fig, use_container_width=True)

            # Pie Chart / Skills distribution
            st.markdown("### 🗂️ Extracted Skills Breakdown")
            skills_count = len(res["skills_extracted"])
            missing_count = len(res["missing_skills"])
            
            fig_pie = go.Figure(data=[go.Pie(
                labels=['Matched Skills', 'Missing Skills'],
                values=[skills_count, missing_count],
                hole=.4,
                marker_colors=['#6366f1', '#f43f5e']
            )])
            fig_pie.update_layout(
                margin=dict(l=10, r=10, t=10, b=10),
                height=200,
                paper_bgcolor='rgba(0,0,0,0)'
            )
            st.plotly_chart(fig_pie, use_container_width=True)

# ── Tab 2: Score Details ────────────────────────────────────
with tab_details:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to inspect score details.")
    else:
        res = st.session_state.analyses[selected_file]
        
        st.markdown(f"### 📉 ATS Compatibility Report: `{selected_file}`")
        st.markdown(f"<div class='feature-box'>{res['score_explanation']}</div>", unsafe_allow_html=True)
        
        col_c1, col_c2 = st.columns([2, 1])
        with col_c1:
            st.markdown("#### Category Analysis")
            df = pd.DataFrame({
                "Category": [c.title() for c in res["score_breakdown"].keys()],
                "Score": list(res["score_breakdown"].values())
            })
            st.dataframe(df, use_container_width=True, hide_index=True)
        with col_c2:
            st.markdown("#### Interview Readiness Rating")
            st.markdown(f"""
            <div style="background: #e0f2fe; border: 1px solid #bae6fd; border-radius: 16px; padding: 1.5rem; text-align: center;">
                <h3 style="color: #0369a1; margin-bottom: 0.5rem;">Readiness Grade</h3>
                <p style="font-size: 1.1rem; color: #075985; line-height: 1.5;">{res['interview_readiness']}</p>
            </div>
            """, unsafe_allow_html=True)

# ── Tab 3: Keywords Matching ────────────────────────────────
with tab_keywords:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to inspect keywords.")
    else:
        res = st.session_state.analyses[selected_file]
        
        col_k1, col_k2 = st.columns(2)
        with col_k1:
            st.markdown("### 🧩 Matched Skills & Keywords")
            st.markdown("These keywords were successfully matched against target standards:")
            if res["skills_extracted"]:
                for s in res["skills_extracted"]:
                    st.markdown(f"⭐ **{s}**")
            else:
                st.write("No specific keywords detected.")
                
        with col_k2:
            st.markdown("### ❌ Missing Skills & Gaps")
            st.markdown("Integrate these skills into your resume to bypass recruiter search rules:")
            for m in res["missing_skills"]:
                st.markdown(f"<span style='background: #ffe4e6; color: #b91c1c; padding: 0.2rem 0.6rem; border-radius: 8px; font-weight: 600; display: inline-block; margin: 0.2rem 0;'>{m}</span>", unsafe_allow_html=True)
                
            st.markdown("<div style='height: 20px;'></div>", unsafe_allow_html=True)
            st.markdown("### 📜 Missing Certifications")
            for cert in res.get("missing_certifications", []):
                st.markdown(f"🎖️ {cert}")

# ── Tab 4: Resume Rewriter ──────────────────────────────────
with tab_rewriter:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to view rewrites.")
    else:
        res = st.session_state.analyses[selected_file]
        
        st.markdown(f"### ✨ AI ATS-Optimized Rewrites: `{selected_file}`")
        st.markdown("Copy and use these improved sections to upgrade your resume layout and key search density.")
        
        # Summary Section
        st.markdown("#### 1. Professional Summary")
        st.markdown(f"<div class='feature-box' style='background: #fdf4ff; border-left-color: #d946ef;'>{res['rewritten_summary']}</div>", unsafe_allow_html=True)
        
        # Skills Section
        st.markdown("#### 2. Categorized Skills Section")
        st.code(res["rewritten_skills"], language="text")
        
        # Experience Section
        st.markdown("#### 3. Optimized Experience Bullet Points")
        for bullet in res["rewritten_experience"]:
            st.markdown(f"<div class='interactive-bullet' style='border-left: 5px solid #10b981;'>🎯 {bullet}</div>", unsafe_allow_html=True)
            
        # Projects Section
        st.markdown("#### 4. Project Highlights")
        for proj in res["rewritten_projects"]:
            st.markdown(f"<div class='interactive-bullet' style='border-left: 5px solid #3b82f6;'>🚀 {proj}</div>", unsafe_allow_html=True)
            
        # Download button
        report_text = f"""RESUME UPGRADE REPORT: {target_role} ({selected_file})
------------------------------------------------
ATS Score: {res['ats_score']}%

RECOMMENDED SUMMARY:
{res['rewritten_summary']}

OPTIMIZED SKILLS:
{res['rewritten_skills']}

EXPERIENCE REWRITES:
{chr(10).join(['- ' + b for b in res['rewritten_experience']])}
"""
        st.download_button(
            f"⬇️ Download Upgrade Guide ({selected_file})", 
            report_text, 
            file_name=f"resume_upgrades_{selected_file.replace('.','_')}.txt"
        )

# ── Tab 5: Career Roadmaps ──────────────────────────────────
with tab_coach:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to view roadmaps.")
    else:
        res = st.session_state.analyses[selected_file]
        coach = res["career_suggestions"]
        
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            st.markdown("### 🎯 Career Development Roles")
            st.markdown("Based on your existing skills and profile, here are other target growth roles:")
            for role in coach["suitable_roles"]:
                st.markdown(f"💼 **{role}**")
                
            st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
            st.markdown("### 🎓 Skills to Learn Next")
            for s in coach["skills_to_learn_next"]:
                st.markdown(f"📚 {s}")
                
            st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
            st.markdown("### 🎖️ Target Professional Certifications")
            for cert in coach["certifications_to_pursue"]:
                st.markdown(f"🏆 {cert}")
                
        with col_c2:
            st.markdown("### 📅 30-Day Personalized Improvement Plan")
            
            # Draw Week-by-Week timeline
            weeks = ["week_1", "week_2", "week_3", "week_4"]
            for i, wk in enumerate(weeks):
                wk_title = f"Week {i+1}"
                wk_desc = res["roadmap_30_day"].get(wk, coach["learning_roadmap"][i] if i < len(coach["learning_roadmap"]) else "")
                st.markdown(f"""
                <div class="roadmap-week">
                    <h5 style="color: #8b5cf6; margin-bottom: 0.25rem;">{wk_title}</h5>
                    <p style="color: #475569; margin: 0; font-size: 0.95rem;">{wk_desc}</p>
                </div>
                """, unsafe_allow_html=True)

# ── Tab 6: Interview Practice ───────────────────────────────
with tab_prep:
    if not has_data or not selected_file:
        st.info("💡 Please upload resumes and click 'Process Resumes' to view interview prep.")
    else:
        res = st.session_state.analyses[selected_file]
        prep = res["interview_prep"]
        
        st.markdown(f"### 💬 Mock Interview Board: `{selected_file}`")
        st.markdown("Test yourself against these custom questions generated based on your resume and target role:")
        
        # Technical Questions
        st.markdown("#### 💻 Technical Questions")
        for i, q in enumerate(prep["technical_questions"]):
            with st.expander(f"Question {i+1}: {q}"):
                st.write("*Tip: Explain your conceptual knowledge first, then walk through a real-world project example where you applied it.*")
                st.text_area("Your response draft:", key=f"tech_ans_{i}", placeholder="Draft your answer here...")
                
        st.divider()
        
        # HR Questions
        st.markdown("#### 👤 HR / Behavioral Questions")
        for i, q in enumerate(prep["hr_questions"]):
            with st.expander(f"Question {i+1}: {q}"):
                st.write("*Tip: Use the STAR framework (Situation, Task, Action, Result) to structure your response.*")
                st.text_area("Your response draft:", key=f"hr_ans_{i}", placeholder="Draft your answer here...")

# ── Tab 7: History ──────────────────────────────────────────
with tab_history:
    st.markdown("### 📜 Past Analysis Runs")
    history_data = get_history()
    
    if not history_data:
        st.info("No analysis records found in SQLite history. Upload a resume to create database entries.")
    else:
        df_list = []
        for row in history_data:
            df_list.append({
                "ID": row[0],
                "Filename": row[1],
                "Timestamp": row[2],
                "ATS Score": f"{row[3]}%",
                "Target Role": row[4]
            })
            
        df_hist = pd.DataFrame(df_list)
        st.dataframe(df_hist, use_container_width=True, hide_index=True)
        
        col_clear, _ = st.columns([1, 4])
        with col_clear:
            if st.button("🗑️ Clear History", use_container_width=True):
                clear_history()
                st.success("History records deleted.")
                st.rerun()
